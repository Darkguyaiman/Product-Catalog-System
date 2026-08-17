from __future__ import annotations

import hashlib
import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def run_props(run):
    rpr = run._element.rPr
    fonts = {}
    if rpr is not None and rpr.rFonts is not None:
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            value = rpr.rFonts.get(qn(f"w:{key}"))
            if value:
                fonts[key] = value
    return {
        "text": run.text,
        "bold": run.bold,
        "italic": run.italic,
        "underline": bool(run.underline) if run.underline is not None else None,
        "size_pt": run.font.size.pt if run.font.size else None,
        "font": run.font.name,
        "fonts_xml": fonts,
        "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
    }


def para_props(p):
    pf = p.paragraph_format
    return {
        "style": p.style.name if p.style else None,
        "alignment": str(p.alignment),
        "left_indent_in": pf.left_indent.inches if pf.left_indent else None,
        "first_line_indent_in": pf.first_line_indent.inches if pf.first_line_indent else None,
        "space_before_pt": pf.space_before.pt if pf.space_before else None,
        "space_after_pt": pf.space_after.pt if pf.space_after else None,
        "line_spacing": str(pf.line_spacing),
        "keep_with_next": pf.keep_with_next,
        "page_break_before": pf.page_break_before,
        "runs": [run_props(r) for r in p.runs],
    }


def main(path_str: str, output_str: str | None = None):
    path = Path(path_str).resolve()
    doc = Document(path)
    data = {
        "path": str(path),
        "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "footers": [],
        "package_parts": [],
    }
    for i, p in enumerate(doc.paragraphs):
        item = para_props(p)
        item["index"] = i
        item["text"] = p.text
        data["paragraphs"].append(item)
    for ti, table in enumerate(doc.tables):
        rows = []
        for ri, row in enumerate(table.rows):
            cells = []
            for ci, cell in enumerate(row.cells):
                cells.append({
                    "index": ci,
                    "text": cell.text,
                    "paragraphs": [dict(para_props(p), text=p.text) for p in cell.paragraphs],
                })
            rows.append({"index": ri, "cells": cells})
        data["tables"].append({"index": ti, "rows": rows})
    for si, sec in enumerate(doc.sections):
        data["headers"].append({"section": si, "paragraphs": [dict(para_props(p), text=p.text) for p in sec.header.paragraphs]})
        data["footers"].append({"section": si, "paragraphs": [dict(para_props(p), text=p.text) for p in sec.footer.paragraphs]})
    with zipfile.ZipFile(path) as zf:
        for info in zf.infolist():
            data["package_parts"].append({"path": info.filename, "size": info.file_size, "sha256": hashlib.sha256(zf.read(info.filename)).hexdigest()})
    payload = json.dumps(data, indent=2, ensure_ascii=False)
    if output_str:
        Path(output_str).write_text(payload, encoding="utf-8")
    else:
        sys.stdout.write(payload)


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None)
