from __future__ import annotations

import hashlib
import math
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "Product Catalog System Documentation"
OUT.mkdir(parents=True, exist_ok=True)

FONT = "Times New Roman"
NAVY = "0F172A"
SLATE = "475569"
HEADER_FILL = "DDE6F0"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=100, bottom=70, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="000000", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)


def set_table_geometry(table, widths_in):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[i])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths_in[i] * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_in) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_run_font(p.add_run(h), 8.5, True)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run_font(p.add_run(value), 8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = False
    set_run_font(p.add_run(text), {1: 18, 2: 14, 3: 12}[level], True)
    return p


def add_body(doc, text, bold_label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    if bold_label and text.startswith(bold_label):
        set_run_font(p.add_run(bold_label), 11, True)
        set_run_font(p.add_run(text[len(bold_label):]), 11)
    else:
        set_run_font(p.add_run(text), 11)
    return p


def add_list(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    num_id = None
    if numbered:
        style_num_id = doc.styles[style]._element.pPr.numPr.numId.val
        numbering = doc.part.numbering_part.element
        base_num = next((n for n in numbering.findall(qn("w:num")) if n.get(qn("w:numId")) == str(style_num_id)), None)
        if base_num is None:
            raise RuntimeError("Could not locate the List Number definition")
        abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
        existing_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
        num_id = max(existing_ids, default=0) + 1
        new_num = OxmlElement("w:num")
        new_num.set(qn("w:numId"), str(num_id))
        abstract = OxmlElement("w:abstractNumId")
        abstract.set(qn("w:val"), abstract_id)
        new_num.append(abstract)
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:startOverride")
        start.set(qn("w:val"), "1")
        override.append(start)
        new_num.append(override)
        numbering.append(new_num)
    for item in items:
        p = doc.add_paragraph(style=style)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        if numbered:
            num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
            num_pr.get_or_add_ilvl().val = 0
            num_pr.get_or_add_numId().val = num_id
        set_run_font(p.add_run(item), 11)


def add_procedure(doc, title, steps):
    add_heading(doc, title, 3)
    add_list(doc, steps, numbered=True)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, 9)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update field to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, placeholder, fld_end])
    set_run_font(run, 11)


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.4)
    sec.footer_distance = Inches(0.45)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    for level, size in ((1, 18), (2, 14), (3, 12)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True
    for sname, size, left in (("TOC 1", 11, 0), ("TOC 2", 10.5, 0.2), ("TOC 3", 10, 0.4)):
        try:
            style = doc.styles[sname]
        except KeyError:
            style = doc.styles.add_style(sname, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.paragraph_format.left_indent = Inches(left)
        style.paragraph_format.space_after = Pt(2)
    footer = sec.footer
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)
    ft = footer.add_table(rows=1, cols=3, width=Inches(6.5))
    ft.autofit = False
    widths = [1.2, 4.1, 1.2]
    for i, cell in enumerate(ft.rows[0].cells):
        cell.width = Inches(widths[i])
        set_cell_margins(cell, 0, 0, 0, 0)
    center_p = ft.cell(0, 1).paragraphs[0]
    center_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(center_p.add_run("Product Catalog System User Guide"), 9, color=SLATE)
    page_p = ft.cell(0, 2).paragraphs[0]
    page_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(page_p)
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def build_docx():
    doc = Document()
    configure_document(doc)

    for _ in range(7):
        doc.add_paragraph()
    logo = ROOT / "public" / "QSS Healthcare.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(logo), width=Inches(4.2))
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run("Product Catalog System"), 24, True)
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("Table of Contents"), 18, True)
    toc_p = doc.add_paragraph()
    add_toc_field(toc_p)
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(p.add_run("Tab Explanation and Step-by-Step User Guide"), 14)
    add_body(doc, "This guide explains each tab in the Product Catalog System, the information shown there, who normally uses it, and the exact steps for common work. It covers the authenticated admin panel and the company-scoped public catalog.")
    add_body(doc, "Reference style: module purpose, information displayed, main actions, workflow, and controls. Role notes below describe the access checks currently implemented by the application.")

    add_heading(doc, "System Flow", 1)
    add_body(doc, "End-to-end process: Sign in -> Review Dashboard -> Maintain reference data -> Add suppliers and companies -> Create products -> Link marketing assets -> Build packages -> Publish active products through each company catalog -> Review and export catalog metrics.", "End-to-end process:")

    add_heading(doc, "Sidebar Structure Overview", 1)
    add_body(doc, "This is the practical navigation map for the authenticated left sidebar. Marketing and Product Catalogue sections can be collapsed or expanded, and their state is remembered by the browser.")
    tree = [
        "Dashboard",
        "Marketing & Materials: Brochures, Fliers, Roll-ups, Posters, Backdrops, Events, Testimonies",
        "Product Catalogue: Packages, Products, Suppliers, Companies, Bulk Import",
        "Settings: Countries, Product Types, Categories, Users",
        "Sign Out",
        "Public company catalog: Home / Products, Packages, Product Detail, Package Detail",
    ]
    add_list(doc, tree)
    add_table(doc, ["Sidebar tab", "Sub pages/actions", "Access note"], [
        ("Dashboard", "Date filter, KPIs, catalog health, charts, PDF download.", "All authenticated roles."),
        ("Marketing & Materials", "Five asset categories plus Events and Testimonies; add, edit, link products, upload files.", "All authenticated roles; delete is Super Admin only."),
        ("Packages", "Search, create, edit, order products, specifications, image upload.", "All authenticated roles; delete is Super Admin only."),
        ("Products", "Search/filter, create, edit, activate/deactivate, images, certificate, specifications.", "All authenticated roles; delete is Super Admin only."),
        ("Suppliers", "Add/edit supplier, country, affiliated-company links.", "All authenticated roles; delete is Super Admin only."),
        ("Companies", "Add/edit company branding, shortname and contact details.", "All authenticated roles; delete is Super Admin only."),
        ("Bulk Import", "Import Countries, Product Types, or Categories from a spreadsheet.", "All authenticated roles."),
        ("Settings", "Countries, Product Types, Categories, and Users.", "Reference data for all signed-in roles; Users only Admin/Super Admin."),
        ("Public Catalog", "Company-specific products, packages, details, downloads and media links.", "No sign-in; active products only, scoped by company shortname."),
    ], [1.45, 3.55, 1.5])

    add_heading(doc, "Role Access", 1)
    add_body(doc, "The routes require a signed-in user for every admin module. Destructive actions and user administration apply additional role checks. Product Specialist and Graphic Designer have broad create/edit access in the current implementation; their names describe operating focus, not a strict route boundary.")
    add_table(doc, ["Role", "Can view", "Can change"], [
        ("Super Admin", "All admin and public-catalog areas.", "Create/edit all records, manage every user, activate/deactivate products, and delete records."),
        ("Admin", "All admin and public-catalog areas.", "Create/edit operational records; manage non-Super-Admin users; cannot use Super-Admin-only delete actions."),
        ("Product Specialist", "All authenticated operational modules except User Management.", "Current routes allow create/edit across catalog, marketing, import, and reference data; no delete or user management."),
        ("Graphic Designer", "All authenticated operational modules except User Management.", "Current routes allow create/edit across catalog, marketing, import, and reference data; no delete or user management."),
        ("Public visitor", "Company-scoped active products and packages.", "No admin changes; browse, search, filter, view media, and open certificates."),
    ], [1.25, 2.55, 2.7])

    add_heading(doc, "1. Dashboard", 1)
    add_body(doc, "Purpose: The Dashboard is the catalog health and activity overview. It summarizes core record counts, product completeness, publishing status, marketing activity, and category/supplier distribution.", "Purpose:")
    add_heading(doc, "Information Displayed", 2)
    add_list(doc, [
        "Total Products, Companies, and Suppliers for the selected date range.",
        "Active and inactive product counts; inactive products do not appear in the public catalog.",
        "Product completeness indicators, including records with images, certificates, categories, suppliers, types, descriptions, and specifications.",
        "Testimonials, Events, and Marketing Assets with change percentages against the previous equivalent period.",
        "Products per Category and Top Suppliers charts.",
    ])
    add_procedure(doc, "How To Review Catalog Health", [
        "Open Dashboard after signing in.",
        "Check Active and Inactive Products to confirm the intended public visibility.",
        "Review completeness indicators and open Products when images, certificates, categories, or other product data are missing.",
        "Use Products per Category to spot over- or under-represented categories.",
        "Use Top Suppliers to review the product distribution by supplier.",
    ])
    add_procedure(doc, "How To Change The Reporting Period", [
        "Open the Date range control.",
        "Choose a preset or enter a custom start and end date.",
        "Click Filter to reload counts, comparisons, and charts.",
        "Click Download PDF to export the same period as a dashboard report.",
    ])

    add_heading(doc, "2. Marketing & Materials", 1)
    add_body(doc, "Purpose: Marketing & Materials stores files and external media references that support product presentation. The sidebar separates Brochures, Fliers, Roll-ups, Posters, Backdrops, Events, and Testimonies.", "Purpose:")
    add_table(doc, ["Sub-tab", "Information", "Main actions"], [
        ("Brochures", "Name, affiliated company, file, linked products.", "Add/edit a company-specific brochure and connect it to products."),
        ("Fliers / Roll-ups / Posters / Backdrops", "Name, file type, linked products.", "Upload an image, PDF, PPTX, or XLSX and link one or more products."),
        ("Events", "Event name, location, start/end dates, media links, linked products.", "Add/edit the event and maintain titled photo/video URLs."),
        ("Testimonies", "Client, treatment, location, treatment date, links, linked products.", "Add/edit a client story and titled review/video URLs."),
    ], [1.4, 2.55, 2.55])
    add_procedure(doc, "How To Add A Marketing Material", [
        "Open the required material category from Marketing & Materials.",
        "Click Add and enter the Title/Name.",
        "For a brochure, select the Affiliated Company so the public catalog can show the correct company-specific file.",
        "Upload the required file. Large files are sent in chunks and assembled by the server.",
        "Select the linked products and save.",
    ])
    add_procedure(doc, "How To Add An Event Or Testimony", [
        "Open Events or Testimonies and click Add.",
        "Complete the identity, location, and date fields.",
        "Add one or more titled media links using full URLs.",
        "Select the products that should display the item.",
        "Save and verify the linked item from a public product detail page.",
    ])
    add_heading(doc, "Deletion Control", 2)
    add_body(doc, "Only Super Admin can delete materials, events, or testimonies. Other authenticated roles can create and edit them. Deleting a stored material can also remove its uploaded file from the server.")

    add_heading(doc, "3. Packages", 1)
    add_body(doc, "Purpose: Packages bundle products into a named solution that can be displayed in the company-specific public catalog.", "Purpose:")
    add_heading(doc, "Information Displayed", 2)
    add_list(doc, ["Package image, name, description, and product count.", "Included products in a deliberate display order.", "Package specifications as name/value pairs with their own order.", "Search and edit actions; delete appears only to Super Admin."])
    add_procedure(doc, "How To Create A Package", [
        "Open Product Catalogue > Packages and click Create Package.",
        "Enter the package name and description.",
        "Upload the main package image.",
        "Choose the included products and arrange their display order.",
        "Add specification name/value pairs and order them for display.",
        "Save the package, then open a company public catalog that contains at least one included active product to verify visibility.",
    ])
    add_heading(doc, "Public Visibility Logic", 2)
    add_body(doc, "A package is returned for a company only when it contains at least one active product supplied by a supplier linked to that company. The package detail page applies the same company and active-product rules.")

    add_heading(doc, "4. Products", 1)
    add_body(doc, "Purpose: Products is the central catalog-maintenance tab. It combines commercial identity, regulatory data, taxonomy, specifications, images, supplier ownership, and publishing status.", "Purpose:")
    add_table(doc, ["Area", "Fields/actions"], [
        ("Identity", "Product Code (required), Model, Description, MDA Registration Number."),
        ("Classification", "Supplier, multiple Product Types, multiple Categories including nested categories."),
        ("Media & compliance", "Multiple product images, one main image, MDA certificate upload."),
        ("Specifications", "Repeatable specification key/value pairs; spreadsheet-assisted entry is supported by the interface."),
        ("Publishing", "Active/Inactive toggle; inactive records are hidden from public product and package views."),
    ], [1.55, 4.95])
    add_procedure(doc, "How To Find Products", [
        "Open Product Catalogue > Products.",
        "Search by product text and use Status to focus on active or inactive records.",
        "Use the multi-select filters for Category, Supplier, and Product Type.",
        "Clear filters before starting a new search when results appear incomplete.",
    ])
    add_procedure(doc, "How To Create A Product", [
        "Click Add Product and upload one or more product images.",
        "Choose the main image that represents the product in lists and public cards.",
        "Enter the required Product Code, then complete Model, MDA registration number, Supplier, Types, Categories, and Description.",
        "Upload the MDA certificate when available.",
        "Add specification name/value pairs and confirm the Active status.",
        "Save, then review the product detail page and linked marketing content.",
    ])
    add_procedure(doc, "How To Edit Visibility Or Delete", [
        "Open Edit for the product and change Product Status to active or inactive as required.",
        "Use inactive status when a product should be retained internally but hidden from every public company catalog.",
        "Only Super Admin can permanently delete a product. Review package and marketing links before deletion.",
    ])

    add_heading(doc, "5. Suppliers", 1)
    add_body(doc, "Purpose: Suppliers identifies product sources and controls which affiliated companies can expose those suppliers' products in their public catalogs.", "Purpose:")
    add_heading(doc, "Information Displayed", 2)
    add_list(doc, ["Supplier name and associated country.", "Linked affiliated companies.", "Products supplied by the record.", "Add, edit, and Super-Admin-only delete actions."])
    add_procedure(doc, "How To Add Or Edit A Supplier", [
        "Open Product Catalogue > Suppliers and click Add Supplier or Edit.",
        "Enter the supplier name.",
        "Select the country from Settings reference data.",
        "Select every affiliated company that should be allowed to display the supplier's products.",
        "Save and verify that linked active products appear under the expected company shortname.",
    ])

    add_heading(doc, "6. Companies", 1)
    add_body(doc, "Purpose: Companies defines each public-catalog tenant, including its URL shortname, branding, registration information, and contact details.", "Purpose:")
    add_heading(doc, "Information Displayed", 2)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.line_spacing = Pt(1)
    spacer.paragraph_format.space_after = Pt(0)
    set_run_font(spacer.add_run(" "), 1)
    add_list(doc, ["Company name and unique shortname used in URLs.", "Logo and branding shown throughout that company's public catalog.", "Registration number, address, email, and contact number.", "Supplier links determine the product population for the company."])
    add_procedure(doc, "How To Add A Company", [
        "Open Product Catalogue > Companies and click Add Company.",
        "Enter the company name and a concise unique shortname.",
        "Upload the company logo and complete registration/contact information.",
        "Save the company.",
        "Link suppliers to the new company from Suppliers, then test /<shortname>/home.",
    ])
    add_heading(doc, "Shortname Control", 2)
    add_body(doc, "The shortname is part of every public route. Changing it changes the catalog URL. Confirm any published links before editing a live shortname.")

    add_heading(doc, "7. Bulk Import", 1)
    add_body(doc, "Purpose: Bulk Import loads reference lists in batches from Excel or CSV files. It supports Countries, Product Types, and Categories.", "Purpose:")
    add_procedure(doc, "How To Import Reference Data", [
        "Open Product Catalogue > Bulk Import.",
        "Choose Countries, Product Types, or Categories.",
        "Open the provided template and keep its expected column headings.",
        "Select an .xlsx, .xls, or .csv file and review the staged filename.",
        "Click Start Import.",
        "Review the import results for inserted, skipped, or failed rows before leaving the page.",
    ])
    add_heading(doc, "Import Controls", 2)
    add_list(doc, ["Use the template for predictable column mapping.", "Import categories in a structure that preserves main/sub-category relationships.", "Verify imported values in Settings before creating products that depend on them."])

    add_heading(doc, "8. Settings", 1)
    add_body(doc, "Purpose: Settings maintains reusable reference data for suppliers and products. Its tabs are Countries, Product Types, Categories, and Users.", "Purpose:")
    add_table(doc, ["Settings tab", "Purpose", "Important actions"], [
        ("Countries", "Country options used by supplier records.", "Add or edit for any signed-in role; delete for Super Admin."),
        ("Product Types", "Reusable type labels assigned to products.", "Add or edit for any signed-in role; delete for Super Admin."),
        ("Categories", "Hierarchical catalog navigation and product classification.", "Add/edit a main or sub-category; delete for Super Admin."),
        ("Users", "Authenticated staff accounts and roles.", "Admin/Super Admin only; delete for Super Admin."),
    ], [1.35, 2.45, 2.7])
    add_procedure(doc, "How To Maintain Countries Or Product Types", [
        "Open Settings and select the required tab.",
        "Enter a new value and add it, or open Edit for an existing value.",
        "Save the change and verify it appears in supplier or product forms.",
        "Super Admin may delete a value after checking whether existing records depend on it.",
    ])
    add_procedure(doc, "How To Maintain Categories", [
        "Open Settings > Categories.",
        "Enter the Category Name.",
        "Leave Main Category empty for a top-level category, or choose a parent for a sub-category.",
        "Save and review the relationship tree.",
        "Verify the category in the public filter after assigning active products to it.",
    ])

    add_heading(doc, "9. User Management", 1)
    add_body(doc, "Purpose: User Management controls who can sign in and which role is stored in the session. The page is available only to Admin and Super Admin.", "Purpose:")
    add_heading(doc, "Information Displayed", 2)
    add_list(doc, ["User email and assigned role.", "Add and edit controls for Admin and Super Admin.", "Delete control for Super Admin only.", "Admin cannot view, create, edit, or promote a Super Admin account."])
    add_procedure(doc, "How To Add Or Edit A User", [
        "Open Settings > Users.",
        "Enter the email, a password of at least six characters, and the role.",
        "Click Add User.",
        "To edit, choose Edit and change the email or role; leave New Password blank to keep the current password.",
        "Only Super Admin can assign the Super Admin role.",
    ])
    add_heading(doc, "Deletion Controls", 2)
    add_body(doc, "Only Super Admin can delete a user. The signed-in Super Admin cannot delete their own account from the user list.")

    add_heading(doc, "10. Public Company Catalog", 1)
    add_body(doc, "Purpose: The public catalog gives each affiliated company a branded, unauthenticated product and package experience. Every query is scoped by the company shortname and supplier-company links.", "Purpose:")
    add_table(doc, ["Public page", "What the visitor can do"], [
        ("Home / Products", "Search active products, filter by hierarchical category, and open a product."),
        ("Product Detail", "Review images, categories, supplier, description, specifications, certificate, marketing materials, events, testimonies, and company brochure."),
        ("Packages", "Search packages available through the company's active product set."),
        ("Package Detail", "Review package image, description, specifications, and included active company-eligible products."),
        ("Watch Video", "Open supported YouTube, Vimeo, event, or testimony media links in an embedded page."),
    ], [1.55, 4.95])
    add_procedure(doc, "How To Validate A Company Catalog", [
        "Open /<shortname>/home without signing in.",
        "Confirm the company name, logo, and contact branding.",
        "Search for a known active product from a supplier linked to the company.",
        "Open the product and verify specifications, images, MDA certificate, and marketing links.",
        "Open Packages and verify that only eligible active products are included.",
    ])
    add_heading(doc, "Public Scope Controls", 2)
    add_list(doc, ["Unknown company shortnames return a not-found page.", "Inactive products are excluded.", "Supplier-company links define company eligibility.", "The company-specific brochure is selected by both product and company."])

    add_heading(doc, "11. Authentication And Sign Out", 1)
    add_body(doc, "Purpose: Authentication protects the admin panel with an email/password login and a database-backed session. Public company catalogs remain accessible without a session.", "Purpose:")
    add_procedure(doc, "How To Sign In And Out", [
        "Open /login or /auth/login.",
        "Enter the registered email and password.",
        "After successful verification, the system stores the user ID, email, and role in the session and opens Dashboard. When finished, click Sign Out in the sidebar footer to destroy the session.",
    ])

    add_heading(doc, "System Controls And Good Practice", 1)
    add_heading(doc, "Required Fields And Validation", 2)
    add_list(doc, [
        "Product Code is required; duplicate or incomplete identifiers should be resolved before publishing.",
        "Material names and uploaded files are required for new marketing materials.",
        "Company shortnames must remain unique because they control public routes.",
        "User emails, passwords, and roles are validated before account creation or editing.",
        "File uploads should use supported image, PDF, office-document, or spreadsheet formats shown by the form.",
    ])
    add_heading(doc, "Recommended Operating Routine", 2)
    add_list(doc, [
        "Maintain Countries, Product Types, and Categories before adding dependent suppliers or products.",
        "Create companies and suppliers, then link them before testing a public catalog.",
        "Create products with a main image, supplier, category, type, specifications, and certificate where applicable.",
        "Add company-specific brochures and supporting marketing links after the product exists.",
        "Build packages only from complete, active products.",
        "Review Dashboard completeness and public catalogs before sharing a shortname URL.",
        "Prefer deactivation over deletion when a product may be needed for history or later republishing.",
    ], numbered=True)
    add_heading(doc, "Complete Information Flow", 2)
    add_body(doc, "Catalog flow: Settings reference data -> Affiliated company -> Supplier and company links -> Product identity, media, compliance, and taxonomy -> Marketing relationships -> Package composition -> Active publishing status -> Company-scoped public catalog -> Dashboard reporting.", "Catalog flow:")

    out_path = OUT / "Product Catalog System Per Tab Explanation.docx"
    doc.save(out_path)
    return out_path


# ---------- Diagram generation ----------

DIAGRAM_COLORS = {
    "navy": ("#0f172a", "#c7d2fe"),
    "hub": ("#0e7490", "#ffffff"),
    "blue": ("#dbeafe", "#1e3a8a"),
    "green": ("#d1fae5", "#065f46"),
    "amber": ("#fef3c7", "#78350f"),
    "pink": ("#fce7f3", "#831843"),
    "purple": ("#ede9fe", "#5b21b6"),
    "slate": ("#f1f5f9", "#334155"),
    "red": ("#fee2e2", "#7f1d1d"),
}


def wrap_text(text, max_chars):
    words = text.split()
    lines, line = [], ""
    for word in words:
        trial = f"{line} {word}".strip()
        if len(trial) <= max_chars or not line:
            line = trial
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


class SvgDiagram:
    def __init__(self, width, height):
        self.w, self.h = width, height
        self.parts = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}" role="graphics-document document">',
                      '<defs><marker id="arrow" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto"><path d="M0,0 L0,6 L9,3 z" fill="#94a3b8"/></marker><filter id="shadow"><feDropShadow dx="7" dy="8" stdDeviation="3" flood-color="#94a3b8" flood-opacity="0.22"/></filter></defs>',
                      '<rect width="100%" height="100%" fill="#ffffff"/>']

    def line(self, x1, y1, x2, y2, dashed=False):
        dash = ' stroke-dasharray="8 7"' if dashed else ""
        self.parts.append(f'<path d="M{x1} {y1} L{x1} {(y1+y2)/2} L{x2} {(y1+y2)/2} L{x2} {y2}" fill="none" stroke="#94a3b8" stroke-width="3" marker-end="url(#arrow)"{dash}/>' )

    def box(self, x, y, w, h, text, kind="slate", title=False, rounded=8, font=28):
        fill, fg = DIAGRAM_COLORS[kind]
        stroke = fg if kind != "navy" else "#0f172a"
        self.parts.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="{rounded}" fill="{fill}" stroke="{stroke}" stroke-width="3" filter="url(#shadow)"/>')
        lines = wrap_text(text, max(12, int(w / (font * 0.6))))
        total = len(lines) * font * 1.28
        start_y = y + (h - total) / 2 + font
        weight = 700 if title else 500
        for i, line in enumerate(lines):
            self.parts.append(f'<text x="{x+w/2}" y="{start_y+i*font*1.28}" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="{font}" font-weight="{weight}" fill="{fg}">{escape(line)}</text>')

    def group(self, x, y, w, h, label, kind):
        fill, fg = DIAGRAM_COLORS[kind]
        self.parts.append(f'<rect x="{x}" y="{y}" width="{w}" height="{h}" rx="16" fill="{fill}" fill-opacity="0.52" stroke="{fg}" stroke-width="3"/>')
        self.parts.append(f'<text x="{x+w/2}" y="{y+38}" text-anchor="middle" font-family="Segoe UI, Arial, sans-serif" font-size="28" font-weight="700" fill="{fg}">{escape(label)}</text>')

    def finish(self, path):
        self.parts.append('</svg>')
        path.write_text("".join(self.parts), encoding="utf-8")


class PdfDiagram:
    def __init__(self, width, height, path):
        self.w, self.h = width, height
        self.c = canvas.Canvas(str(path), pagesize=(width, height), pageCompression=1)
        self.c.setFillColorRGB(1, 1, 1)
        self.c.rect(0, 0, width, height, stroke=0, fill=1)

    @staticmethod
    def hex_color(hex_value):
        h = hex_value.lstrip("#")
        return tuple(int(h[i:i+2], 16) / 255 for i in (0, 2, 4))

    def line(self, x1, y1, x2, y2, dashed=False):
        c = self.c
        c.saveState()
        c.setStrokeColorRGB(*self.hex_color("#94a3b8"))
        c.setFillColorRGB(*self.hex_color("#94a3b8"))
        c.setLineWidth(3)
        if dashed:
            c.setDash(8, 7)
        yy1, yy2 = self.h - y1, self.h - y2
        mid = (yy1 + yy2) / 2
        p = c.beginPath()
        p.moveTo(x1, yy1)
        p.lineTo(x1, mid)
        p.lineTo(x2, mid)
        p.lineTo(x2, yy2 + 8)
        c.drawPath(p)
        c.line(x2, yy2 + 8, x2 - 7, yy2 + 18)
        c.line(x2, yy2 + 8, x2 + 7, yy2 + 18)
        c.restoreState()

    def box(self, x, y, w, h, text, kind="slate", title=False, rounded=8, font=28):
        fill, fg = DIAGRAM_COLORS[kind]
        c = self.c
        c.saveState()
        c.setFillColorRGB(*self.hex_color(fill))
        c.setStrokeColorRGB(*self.hex_color(fg if kind != "navy" else "#0f172a"))
        c.setLineWidth(3)
        c.roundRect(x, self.h-y-h, w, h, rounded, stroke=1, fill=1)
        lines = wrap_text(text, max(12, int(w / (font * 0.6))))
        total = len(lines) * font * 1.28
        baseline = self.h - y - (h - total) / 2 - font
        c.setFillColorRGB(*self.hex_color(fg))
        c.setFont("Helvetica-Bold" if title else "Helvetica", font)
        for i, line in enumerate(lines):
            c.drawCentredString(x+w/2, baseline-i*font*1.28, line)
        c.restoreState()

    def group(self, x, y, w, h, label, kind):
        fill, fg = DIAGRAM_COLORS[kind]
        c = self.c
        c.saveState()
        c.setFillColorRGB(*self.hex_color(fill), alpha=0.52)
        c.setStrokeColorRGB(*self.hex_color(fg))
        c.setLineWidth(3)
        c.roundRect(x, self.h-y-h, w, h, 16, stroke=1, fill=1)
        c.setFillColorRGB(*self.hex_color(fg))
        c.setFont("Helvetica-Bold", 28)
        c.drawCentredString(x+w/2, self.h-y-39, label)
        c.restoreState()

    def finish(self):
        self.c.showPage()
        self.c.save()


def draw_role_diagram(role, groups, unavailable, filename):
    width, height = 4800, 980
    svg_path = OUT / f"{filename}.svg"
    pdf_path = OUT / f"{filename}.pdf"
    renderers = [SvgDiagram(width, height), PdfDiagram(width, height, pdf_path)]
    title = f"{role} Role - How to Use Each Tab"
    for d in renderers:
        d.box(1960, 35, 880, 125, title, "navy", True, 0, 34)
        d.box(2150, 200, 500, 80, f"Sign in as {role}", "hub", False, 40, 28)
    cols = len(groups)
    gap = 35
    group_w = (width - 160 - gap*(cols-1)) / cols
    group_y, group_h = 380, 455
    for idx, (label, kind, steps) in enumerate(groups):
        x = 80 + idx*(group_w+gap)
        for d in renderers:
            d.line(2400, 280, x+group_w/2, group_y)
            d.group(x, group_y, group_w, group_h, label, kind)
        step_gap = 18
        step_w = group_w - 60
        step_h = (group_h - 85 - step_gap*(len(steps)-1)) / len(steps)
        sy = group_y + 60
        for si, step in enumerate(steps):
            for d in renderers:
                d.box(x+30, sy+si*(step_h+step_gap), step_w, step_h, step, "slate", False, 4, 21)
    if unavailable:
        y = 875
        for d in renderers:
            d.box(1250, y, 2300, 70, unavailable, "red", False, 6, 22)
    renderers[0].finish(svg_path)
    renderers[1].finish()
    return svg_path, pdf_path


def draw_overview():
    width, height = 5200, 1500
    svg_path = OUT / "Full System Overview Diagram Product Catalog System.svg"
    pdf_path = OUT / "Full System Overview Diagram Product Catalog System.pdf"
    renderers = [SvgDiagram(width, height), PdfDiagram(width, height, pdf_path)]
    for d in renderers:
        d.box(1800, 35, 1600, 120, "Full System Overview Diagram - Product Catalog System", "navy", True, 0, 34)
        d.box(2400, 205, 400, 80, "Sign in", "hub", False, 40, 30)
        d.box(2180, 345, 840, 110, "Dashboard - counts, health, charts, PDF", "blue", True, 8, 27)
        d.line(2600, 285, 2600, 345)
    modules = [
        ("Marketing & Materials", "Assets, events, testimonies, product links", "pink"),
        ("Products", "Identity, images, MDA, taxonomy, specifications, status", "blue"),
        ("Packages", "Bundle active products and ordered specifications", "purple"),
        ("Suppliers", "Country and affiliated-company relationships", "green"),
        ("Companies", "Shortname, logo, branding and public tenant", "amber"),
        ("Bulk Import", "Countries, product types and categories", "green"),
        ("Settings", "Countries, types, categories and users", "pink"),
        ("Public Catalog", "Company-scoped active products and packages", "amber"),
    ]
    positions = []
    for row in range(2):
        for col in range(4):
            positions.append((130 + col*1270, 585 + row*360))
    for index, ((name, desc, kind), (x, y)) in enumerate(zip(modules, positions)):
        for d in renderers:
            if index < 4:
                d.line(2600, 455, x+540, y)
            else:
                top_x, top_y = positions[index - 4]
                d.line(top_x+540, top_y+270, x+540, y)
            d.group(x, y, 1080, 270, name, kind)
            d.box(x+70, y+75, 940, 155, desc, "slate", False, 6, 24)
    roles = [
        "Super Admin - full access and deletion",
        "Admin - operations and non-Super-Admin users",
        "Product Specialist - operational create/edit, no users/delete",
        "Graphic Designer - operational create/edit, no users/delete",
        "Public visitor - browse company catalog",
    ]
    for d in renderers:
        d.group(1280, 1300, 2640, 145, "Who can do what", "slate")
    x = 1320
    widths = [400, 520, 590, 590, 400]
    kinds = ["pink", "blue", "green", "purple", "amber"]
    for text_value, w, kind in zip(roles, widths, kinds):
        for d in renderers:
            d.box(x, 1350, w, 70, text_value, kind, False, 5, 18)
        x += w + 25
    renderers[0].finish(svg_path)
    renderers[1].finish()
    return svg_path, pdf_path


def build_diagrams():
    role_specs = [
        ("Super Admin", [
            ("Dashboard", "blue", ["Review KPIs and catalog health", "Filter period and download PDF"]),
            ("Marketing", "pink", ["Add/edit assets, events and testimonies", "Delete records when approved"]),
            ("Catalogue", "green", ["Create/edit/delete products and packages", "Manage suppliers and companies"]),
            ("Settings", "purple", ["Maintain countries, types and categories", "Manage every user and role"]),
            ("Public Catalog", "amber", ["Validate company branding and visibility", "Confirm active products and packages"]),
        ], "Full access. Review dependencies before permanent deletion.", "Super Admin Role User Flow Diagram"),
        ("Admin", [
            ("Dashboard", "blue", ["Review KPIs and catalog health", "Filter period and download PDF"]),
            ("Marketing", "pink", ["Add/edit assets, events and testimonies", "Link content to products"]),
            ("Catalogue", "green", ["Create/edit products and packages", "Manage suppliers, companies and imports"]),
            ("Settings", "purple", ["Maintain reference data", "Manage non-Super-Admin users"]),
            ("Public Catalog", "amber", ["Validate company branding", "Confirm active visibility"]),
        ], "Not available: Super-Admin-only deletion or managing Super Admin accounts.", "Admin Role User Flow Diagram"),
        ("Product Specialist", [
            ("Dashboard", "blue", ["Review product health and category charts", "Export the dashboard report"]),
            ("Products", "green", ["Create/edit product data and media", "Activate or deactivate products"]),
            ("Catalog Support", "amber", ["Maintain packages, suppliers and companies", "Import or edit reference data"]),
            ("Marketing", "pink", ["Link supporting assets and stories", "Edit files, events and testimonies"]),
            ("Public Catalog", "purple", ["Verify product and package presentation", "Test search and category filters"]),
        ], "Not available: User Management and all Super-Admin-only delete actions.", "Product Specialist Role User Flow Diagram"),
        ("Graphic Designer", [
            ("Dashboard", "blue", ["Review marketing and completeness counts", "Export the dashboard report"]),
            ("Marketing", "pink", ["Upload brochures, fliers, roll-ups, posters and backdrops", "Maintain events, testimonies and links"]),
            ("Product Media", "green", ["Edit product images and main image", "Review descriptions and certificates"]),
            ("Packages & Branding", "purple", ["Update package images and layout data", "Maintain company logos and branding"]),
            ("Public Catalog", "amber", ["Inspect responsive public presentation", "Verify assets on product detail pages"]),
        ], "Not available: User Management and all Super-Admin-only delete actions.", "Graphic Designer Role User Flow Diagram"),
    ]
    outputs = []
    for role, groups, unavailable, filename in role_specs:
        outputs.extend(draw_role_diagram(role, groups, unavailable, filename))
    outputs.extend(draw_overview())
    return outputs


if __name__ == "__main__":
    docx = build_docx()
    diagrams = build_diagrams()
    source_hash = hashlib.sha256((ROOT / "Example" / "System Per Tab Explaination .docx").read_bytes()).hexdigest()
    if source_hash != "88270f62762c4a7bffaa8ae44cd37c220c4b66b603152c43591d531f7b7106d9":
        raise RuntimeError("Reference DOCX changed during generation")
    print(docx)
    for item in diagrams:
        print(item)
